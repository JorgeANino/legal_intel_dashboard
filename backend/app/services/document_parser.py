"""
Document parser for PDF and DOCX files
"""

# Third-party imports
import docx
from pypdf import PdfReader


class DocumentParser:
    """Parse PDF and DOCX files to extract text"""

    @staticmethod
    def parse_pdf(file_path: str) -> tuple[str, int]:
        """
        Extract text from PDF file

        Returns:
            tuple: (extracted_text, page_count)
        """
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PdfReader(file)
                page_count = len(pdf_reader.pages)

                text_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                full_text = "\n\n".join(text_parts)
                return full_text, page_count

        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> tuple[str, int]:
        """
        Extract text from DOCX file

        Returns:
            tuple: (extracted_text, page_count)
        """
        try:
            doc = docx.Document(file_path)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n\n".join(text_parts)

            # Approximate page count (assuming ~500 words per page)
            word_count = len(full_text.split())
            page_count = max(1, word_count // 500)

            return full_text, page_count

        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")

    @classmethod
    def parse_document(cls, file_path: str, file_type: str) -> tuple[str, int]:
        """
        Parse document based on file type

        Args:
            file_path: Path to the document
            file_type: File extension (pdf or docx)

        Returns:
            tuple: (extracted_text, page_count)
        """
        if file_type == "pdf":
            return cls.parse_pdf(file_path)
        elif file_type == "docx":
            return cls.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
