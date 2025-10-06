"""
Generate test documents in PDF and DOCX formats
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def create_nda_docx():
    """Create NDA document in DOCX format"""
    doc = Document()
    
    # Title
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content
    doc.add_paragraph(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of January 15, 2024, '
        'by and between:'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('PARTY A: ').bold = True
    p.add_run(
        'TechVision Solutions LLC, a technology company incorporated under the laws of the '
        'United Arab Emirates, having its principal office at Dubai Technology Park, Dubai, '
        'UAE ("Disclosing Party")'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('AND').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('PARTY B: ').bold = True
    p.add_run(
        'DataSecure Middle East FZ-LLC, a software development company incorporated in Abu Dhabi, '
        'UAE, having its principal office at Masdar City, Abu Dhabi, UAE ("Receiving Party")'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('WHEREAS, ').bold = True
    p.add_run(
        'the Disclosing Party possesses certain confidential and proprietary information related '
        'to artificial intelligence and machine learning technologies;'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('WHEREAS, ').bold = True
    p.add_run(
        'the Receiving Party desires to receive such confidential information for the purpose of '
        'evaluating a potential business collaboration in the field of technology and software development;'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('NOW, THEREFORE, ').bold = True
    p.add_run(
        'in consideration of the mutual covenants and agreements contained herein, the parties agree as follows:'
    )
    
    # Sections
    doc.add_heading('1. CONFIDENTIAL INFORMATION', 2)
    doc.add_paragraph(
        'For purposes of this Agreement, "Confidential Information" means all technical, business, '
        'financial, and other information disclosed by the Disclosing Party, including but not limited to: '
        'software code, algorithms, business plans, customer lists, pricing information, and proprietary '
        'technology specifications.'
    )
    
    doc.add_heading('2. OBLIGATIONS OF RECEIVING PARTY', 2)
    doc.add_paragraph('The Receiving Party agrees to:')
    doc.add_paragraph('(a) Hold all Confidential Information in strict confidence;', style='List Number')
    doc.add_paragraph(
        '(b) Not disclose any Confidential Information to third parties without prior written consent;',
        style='List Number'
    )
    doc.add_paragraph(
        '(c) Use the Confidential Information solely for the purpose stated herein;',
        style='List Number'
    )
    doc.add_paragraph(
        '(d) Protect the Confidential Information using the same degree of care used for its own '
        'confidential information.',
        style='List Number'
    )
    
    doc.add_heading('3. TERM AND TERMINATION', 2)
    doc.add_paragraph(
        'This Agreement shall commence on the Effective Date and shall continue for a period of three (3) years. '
        'The confidentiality obligations shall survive termination for an additional two (2) years.'
    )
    
    doc.add_heading('4. GOVERNING LAW AND JURISDICTION', 2)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of the United Arab Emirates. '
        'Any disputes arising under this Agreement shall be subject to the exclusive jurisdiction of the courts of '
        'Abu Dhabi, UAE.'
    )
    
    doc.add_heading('5. GENERAL PROVISIONS', 2)
    doc.add_paragraph(
        'This Agreement constitutes the entire agreement between the parties concerning the subject matter hereof '
        'and supersedes all prior agreements and understandings.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('IN WITNESS WHEREOF, ').bold = True
    p.add_run('the parties have executed this Agreement as of the date first written above.')
    
    doc.add_paragraph()
    doc.add_paragraph('TechVision Solutions LLC')
    doc.add_paragraph('Authorized Signatory: Ahmed Al-Mansouri')
    doc.add_paragraph('Date: January 15, 2024')
    
    doc.add_paragraph()
    doc.add_paragraph('DataSecure Middle East FZ-LLC')
    doc.add_paragraph('Authorized Signatory: Sarah Thompson')
    doc.add_paragraph('Date: January 15, 2024')
    
    doc.save('01_nda_abudhabi_tech.docx')
    print("✓ Created: 01_nda_abudhabi_tech.docx")


def create_msa_pdf():
    """Create MSA document in PDF format"""
    filename = '02_msa_london_finance.pdf'
    doc = SimpleDocTemplate(filename, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    Story = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='black',
        spaceAfter=30,
        alignment=1  # Center
    )
    Story.append(Paragraph('MASTER SERVICES AGREEMENT', title_style))
    Story.append(Spacer(1, 12))
    
    # Content
    Story.append(Paragraph(
        'This Master Services Agreement ("Agreement") is made and entered into as of March 22, 2024, '
        'by and between:',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph(
        '<b>CLIENT:</b> FinTech Innovations Ltd, a financial technology company registered in England and Wales '
        'under company number 12345678, with its registered office at 10 Canary Wharf, London E14 5AB, '
        'United Kingdom ("Client")',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>AND</b>', styles['BodyText']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph(
        '<b>SERVICE PROVIDER:</b> CloudServices Europe Limited, a software services company registered in '
        'England and Wales, with offices at Tech Hub, Manchester M1 1AB, United Kingdom ("Provider")',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>RECITALS:</b>', styles['Heading2']))
    Story.append(Paragraph(
        'Client desires to engage Provider to perform certain software development and cloud infrastructure '
        'services in the finance and banking sector, and Provider agrees to provide such services in accordance '
        'with the terms and conditions set forth herein.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>1. SERVICES</b>', styles['Heading2']))
    Story.append(Paragraph('Provider shall provide the following services to Client:', styles['BodyText']))
    Story.append(Paragraph('- Cloud infrastructure management and optimization', styles['BodyText']))
    Story.append(Paragraph('- Software development for financial applications', styles['BodyText']))
    Story.append(Paragraph('- Technical support and maintenance services', styles['BodyText']))
    Story.append(Paragraph('- Security and compliance consulting for financial systems', styles['BodyText']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>2. TERM</b>', styles['Heading2']))
    Story.append(Paragraph(
        'This Agreement shall commence on April 1, 2024 (the "Commencement Date") and shall continue for an '
        'initial term of twenty-four (24) months, unless earlier terminated in accordance with Section 8.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>3. COMPENSATION</b>', styles['Heading2']))
    Story.append(Paragraph(
        '3.1 Fees: Client shall pay Provider a monthly retainer fee of £50,000 (Fifty Thousand British Pounds) '
        'plus applicable VAT.',
        styles['BodyText']
    ))
    Story.append(Paragraph(
        '3.2 Additional Services: Any services beyond the scope defined herein shall be billed at Provider\'s '
        'standard hourly rate of £150 per hour.',
        styles['BodyText']
    ))
    Story.append(Paragraph(
        '3.3 Payment Terms: All invoices shall be paid within thirty (30) days of receipt.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>4. INTELLECTUAL PROPERTY</b>', styles['Heading2']))
    Story.append(Paragraph(
        'All work product, deliverables, and intellectual property created by Provider in the course of '
        'performing services under this Agreement shall be the exclusive property of Client.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>5. CONFIDENTIALITY</b>', styles['Heading2']))
    Story.append(Paragraph(
        'Each party agrees to maintain the confidentiality of the other party\'s proprietary and confidential '
        'information and shall not disclose such information to any third party without prior written consent.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>6. LIABILITY AND INDEMNIFICATION</b>', styles['Heading2']))
    Story.append(Paragraph(
        'Provider\'s total liability under this Agreement shall not exceed the total fees paid by Client in '
        'the twelve (12) months preceding the claim. Provider agrees to indemnify Client against any claims '
        'arising from Provider\'s negligence or breach of this Agreement.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>7. TERMINATION</b>', styles['Heading2']))
    Story.append(Paragraph(
        'Either party may terminate this Agreement upon ninety (90) days written notice. In the event of '
        'material breach, the non-breaching party may terminate immediately upon written notice.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>8. GOVERNING LAW</b>', styles['Heading2']))
    Story.append(Paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of England and Wales. '
        'The parties submit to the exclusive jurisdiction of the courts of England and Wales for all disputes '
        'arising under this Agreement.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('<b>9. GENERAL PROVISIONS</b>', styles['Heading2']))
    Story.append(Paragraph(
        'This Agreement represents the entire understanding between the parties and supersedes all prior '
        'negotiations, representations, or agreements.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 24))
    
    Story.append(Paragraph(
        '<b>IN WITNESS WHEREOF,</b> the parties hereto have executed this Master Services Agreement as of '
        'the date first above written.',
        styles['BodyText']
    ))
    Story.append(Spacer(1, 24))
    
    Story.append(Paragraph('FinTech Innovations Ltd', styles['BodyText']))
    Story.append(Paragraph('By: James Robertson, Chief Executive Officer', styles['BodyText']))
    Story.append(Paragraph('Date: March 22, 2024', styles['BodyText']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph('CloudServices Europe Limited', styles['BodyText']))
    Story.append(Paragraph('By: Emily Watson, Managing Director', styles['BodyText']))
    Story.append(Paragraph('Date: March 22, 2024', styles['BodyText']))
    
    doc.build(Story)
    print("✓ Created: 02_msa_london_finance.pdf")


def create_service_agreement_docx():
    """Create Service Agreement in DOCX format"""
    doc = Document()
    
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'This Professional Services Agreement ("Agreement") is entered into effective as of February 10, 2024, between:'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('COMPANY: ').bold = True
    p.add_run(
        'HealthTech Solutions Inc., a Delaware corporation with principal offices at 1234 Healthcare Drive, '
        'Wilmington, Delaware 19801, United States ("Company")'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('AND').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('CONSULTANT: ').bold = True
    p.add_run(
        'MedData Analytics LLC, a Delaware limited liability company with offices at 5678 Medical Plaza, '
        'Dover, Delaware 19901, United States ("Consultant")'
    )
    
    doc.add_paragraph()
    doc.add_heading('BACKGROUND:', 2)
    doc.add_paragraph(
        'Company operates in the healthcare and medical technology sector and desires to engage Consultant to '
        'provide data analytics and healthcare technology consulting services. Consultant has expertise in '
        'healthcare data systems and agrees to provide such services subject to the terms and conditions of this Agreement.'
    )
    
    doc.add_heading('1. SCOPE OF SERVICES', 2)
    doc.add_paragraph('Consultant shall provide the following professional services to Company:')
    doc.add_paragraph('a) Healthcare data analytics and reporting', style='List Bullet')
    doc.add_paragraph('b) Electronic health record (EHR) system optimization', style='List Bullet')
    doc.add_paragraph('c) HIPAA compliance consulting and audit support', style='List Bullet')
    doc.add_paragraph('d) Medical billing system integration and support', style='List Bullet')
    doc.add_paragraph('e) Healthcare technology strategic planning', style='List Bullet')
    
    doc.add_heading('2. TERM AND RENEWAL', 2)
    doc.add_paragraph(
        'This Agreement shall commence on March 1, 2024 and shall continue for an initial term of twelve (12) months '
        '(the "Initial Term"). Upon expiration of the Initial Term, this Agreement shall automatically renew for '
        'successive one-year terms unless either party provides written notice of non-renewal at least sixty (60) days '
        'prior to the expiration of the then-current term.'
    )
    
    doc.add_heading('3. COMPENSATION AND PAYMENT', 2)
    doc.add_paragraph(
        '3.1 Professional Fees: Company shall pay Consultant a monthly fee of $25,000 USD (Twenty-Five Thousand US Dollars) '
        'for the services rendered under this Agreement.'
    )
    doc.add_paragraph(
        '3.2 Expenses: Company shall reimburse Consultant for all reasonable and pre-approved out-of-pocket expenses '
        'incurred in connection with the performance of services.'
    )
    doc.add_paragraph(
        '3.3 Invoicing: Consultant shall submit monthly invoices, and Company shall pay all undisputed amounts within '
        'thirty (30) days of receipt.'
    )
    
    doc.add_heading('4. INTELLECTUAL PROPERTY RIGHTS', 2)
    doc.add_paragraph(
        'Any and all deliverables, work products, inventions, and intellectual property created by Consultant during the '
        'term of this Agreement shall be considered "work made for hire" and shall be the sole and exclusive property of Company.'
    )
    
    doc.add_heading('5. CONFIDENTIALITY AND DATA PROTECTION', 2)
    doc.add_paragraph(
        'Consultant acknowledges that it will have access to Company\'s confidential information and protected health '
        'information (PHI) as defined under HIPAA. Consultant agrees to:'
    )
    doc.add_paragraph('a) Maintain strict confidentiality of all Company information', style='List Bullet')
    doc.add_paragraph('b) Comply with all applicable HIPAA regulations and requirements', style='List Bullet')
    doc.add_paragraph('c) Implement appropriate safeguards to protect PHI', style='List Bullet')
    doc.add_paragraph('d) Not disclose any confidential information without prior written authorization', style='List Bullet')
    
    doc.add_heading('6. GOVERNING LAW AND DISPUTE RESOLUTION', 2)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware, '
        'without regard to its conflict of laws principles. Any disputes arising out of or relating to this Agreement '
        'shall be resolved through binding arbitration in Delaware in accordance with the rules of the American Arbitration Association.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('IN WITNESS WHEREOF, ').bold = True
    p.add_run('the parties have executed this Agreement as of the date first written above.')
    
    doc.add_paragraph()
    doc.add_paragraph('HEALTHTECH SOLUTIONS INC.')
    doc.add_paragraph('By: Dr. Michael Chen')
    doc.add_paragraph('Title: Chief Executive Officer')
    doc.add_paragraph('Date: February 10, 2024')
    
    doc.add_paragraph()
    doc.add_paragraph('MEDDATA ANALYTICS LLC')
    doc.add_paragraph('By: Jennifer Martinez')
    doc.add_paragraph('Title: Managing Partner')
    doc.add_paragraph('Date: February 10, 2024')
    
    doc.save('03_service_agreement_delaware_healthcare.docx')
    print("✓ Created: 03_service_agreement_delaware_healthcare.docx")


if __name__ == '__main__':
    import os
    os.chdir('/Users/jorgenino/Documents/legal_intel_dashboard/test_documents')
    
    print("Generating test documents...")
    print()
    
    create_nda_docx()
    create_msa_pdf()
    create_service_agreement_docx()
    
    print()
    print("✅ Test documents generated successfully!")
    print()
    print("Documents created:")
    print("  1. 01_nda_abudhabi_tech.docx - UAE Tech NDA")
    print("  2. 02_msa_london_finance.pdf - UK Finance MSA")
    print("  3. 03_service_agreement_delaware_healthcare.docx - Delaware Healthcare Service Agreement")

