"""
Generate additional test documents in PDF and DOCX formats
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


def create_franchise_agreement_pdf():
    """Create Franchise Agreement in PDF format"""
    filename = "04_franchise_agreement_california.pdf"
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    Story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=16,
        textColor="black",
        spaceAfter=30,
        alignment=1,
    )

    Story.append(Paragraph("FRANCHISE AGREEMENT", title_style))
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph(
            'This Franchise Agreement ("Agreement") is made and entered into as of April 5, 2024, by and between:',
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph(
            "<b>FRANCHISOR:</b> QuickBite Restaurant Systems Inc., a California corporation with its principal place of "
            'business at 789 Franchise Boulevard, Los Angeles, California 90001, United States ("Franchisor")',
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>AND</b>", styles["BodyText"]))
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph(
            "<b>FRANCHISEE:</b> Golden Gate Dining LLC, a California limited liability company with its principal place of "
            'business at 456 Market Street, San Francisco, California 94102, United States ("Franchisee")',
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>RECITALS:</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "WHEREAS, Franchisor has developed a distinctive system for establishing and operating fast-casual restaurants "
            'specializing in healthy food options under the trademark "QuickBite" in the restaurant and hospitality industry;',
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph(
            "WHEREAS, Franchisee desires to obtain the right to establish and operate a QuickBite restaurant using the "
            "Franchisor's proprietary system, trademarks, and business methods;",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>1. GRANT OF FRANCHISE</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "Franchisor hereby grants to Franchisee, and Franchisee accepts, a non-exclusive franchise to establish and "
            'operate one (1) QuickBite restaurant (the "Franchised Business") at the following location: 456 Market Street, '
            "San Francisco, California 94102, subject to Franchisor's approval of the specific site.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>2. TERM</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "The term of this Agreement shall be ten (10) years, commencing on the Opening Date, which shall be the date "
            "the Franchised Business opens for business to the public.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>3. FRANCHISE FEE AND ROYALTIES</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "3.1 Initial Franchise Fee: Franchisee shall pay Franchisor an initial franchise fee of $50,000 USD "
            "(Fifty Thousand US Dollars), which shall be due and payable upon execution of this Agreement. This fee is non-refundable.",
            styles["BodyText"],
        )
    )
    Story.append(
        Paragraph(
            "3.2 Continuing Royalty Fee: Franchisee shall pay to Franchisor a continuing royalty fee equal to six percent (6%) "
            "of Gross Sales, payable monthly within ten (10) days after the end of each calendar month.",
            styles["BodyText"],
        )
    )
    Story.append(
        Paragraph(
            "3.3 Marketing Fee: Franchisee shall contribute three percent (3%) of Gross Sales to the national marketing fund, "
            "payable monthly.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph("<b>4. TRADEMARKS AND PROPRIETARY MARKS</b>", styles["Heading2"])
    )
    Story.append(
        Paragraph(
            "Franchisee acknowledges that Franchisor is the owner of the QuickBite trademark and all related proprietary marks. "
            "Franchisee is granted a limited license to use such marks solely in connection with the operation of the Franchised "
            "Business and in accordance with this Agreement.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>5. GOVERNING LAW</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "This Agreement shall be governed by the laws of the State of California. Any disputes shall be resolved through "
            "mediation, and if unsuccessful, through binding arbitration conducted in Los Angeles County, California.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 24))

    Story.append(
        Paragraph(
            "<b>IN WITNESS WHEREOF,</b> the parties have executed this Franchise Agreement as of the date first written above.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 24))

    Story.append(Paragraph("QUICKBITE RESTAURANT SYSTEMS INC.", styles["BodyText"]))
    Story.append(Paragraph("By: Robert Thompson", styles["BodyText"]))
    Story.append(Paragraph("Title: President and CEO", styles["BodyText"]))
    Story.append(Paragraph("Date: April 5, 2024", styles["BodyText"]))
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("GOLDEN GATE DINING LLC", styles["BodyText"]))
    Story.append(Paragraph("By: Lisa Wong", styles["BodyText"]))
    Story.append(Paragraph("Title: Managing Member", styles["BodyText"]))
    Story.append(Paragraph("Date: April 5, 2024", styles["BodyText"]))

    doc.build(Story)
    print("✓ Created: 04_franchise_agreement_california.pdf")


def create_license_agreement_docx():
    """Create License Agreement in DOCX format"""
    doc = Document()

    title = doc.add_heading("SOFTWARE LICENSE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'This Software License Agreement ("Agreement") is entered into as of May 18, 2024, by and between:'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("LICENSOR: ").bold = True
    p.add_run(
        "PetroTech Analytics Corporation, a New York corporation with its principal office at 100 Energy Plaza, "
        'New York, NY 10004, United States ("Licensor")'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("AND").bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("LICENSEE: ").bold = True
    p.add_run(
        "Global Energy Solutions Inc., an international oil and gas company incorporated in Delaware with principal "
        'offices at 200 Petroleum Way, Houston, Texas 77002, United States ("Licensee")'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("WHEREAS, ").bold = True
    p.add_run(
        "Licensor has developed proprietary software for oil and gas exploration, drilling optimization, and reservoir "
        'analysis (the "Software");'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("WHEREAS, ").bold = True
    p.add_run(
        "Licensee desires to obtain a license to use the Software for its oil and gas operations in North America and "
        "the Middle East;"
    )

    doc.add_heading("1. GRANT OF LICENSE", 2)
    doc.add_paragraph(
        "Subject to the terms and conditions of this Agreement, Licensor hereby grants to Licensee a non-exclusive, "
        "non-transferable license to:"
    )
    doc.add_paragraph(
        "a) Install and use the Software on up to fifty (50) workstations",
        style="List Bullet",
    )
    doc.add_paragraph(
        "b) Use the Software solely for Licensee's internal business operations in the oil and gas industry",
        style="List Bullet",
    )
    doc.add_paragraph(
        "c) Access the Software's cloud-based features for data analysis and reporting",
        style="List Bullet",
    )

    doc.add_heading("2. LICENSE TERM", 2)
    doc.add_paragraph(
        'This Agreement shall commence on June 1, 2024 (the "Effective Date") and shall continue for an initial term '
        "of three (3) years. The Agreement shall automatically renew for successive one-year terms unless either party "
        "provides written notice of non-renewal at least ninety (90) days prior to the end of the then-current term."
    )

    doc.add_heading("3. LICENSE FEES AND PAYMENT", 2)
    doc.add_paragraph(
        "3.1 Initial License Fee: Licensee shall pay an initial license fee of $500,000 USD (Five Hundred Thousand US Dollars), "
        "payable upon execution of this Agreement."
    )
    doc.add_paragraph(
        "3.2 Annual Maintenance Fee: Beginning on the first anniversary of the Effective Date, Licensee shall pay an "
        "annual maintenance and support fee of $100,000 USD, which includes software updates, technical support, and "
        "maintenance services."
    )

    doc.add_heading("4. INTELLECTUAL PROPERTY RIGHTS", 2)
    doc.add_paragraph(
        "Licensee acknowledges that the Software and all intellectual property rights therein are and shall remain the "
        "sole and exclusive property of Licensor."
    )

    doc.add_heading("5. DATA SECURITY AND COMPLIANCE", 2)
    doc.add_paragraph(
        "Licensor shall implement industry-standard security measures to protect Licensee's drilling data, exploration "
        "data, and other proprietary information. The Software shall comply with all applicable regulations governing the "
        "oil and gas industry in North America and the Middle East."
    )

    doc.add_heading("6. GOVERNING LAW AND JURISDICTION", 2)
    doc.add_paragraph(
        "This Agreement shall be governed by the laws of the State of New York, without regard to its conflicts of law "
        "principles. The parties consent to the exclusive jurisdiction of the courts located in New York County, New York "
        "for any disputes arising under this Agreement."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("IN WITNESS WHEREOF, ").bold = True
    p.add_run(
        "the parties have executed this Software License Agreement as of the date first written above."
    )

    doc.add_paragraph()
    doc.add_paragraph("PETROTECH ANALYTICS CORPORATION")
    doc.add_paragraph("By: David Richardson")
    doc.add_paragraph("Title: Chief Executive Officer")
    doc.add_paragraph("Date: May 18, 2024")

    doc.add_paragraph()
    doc.add_paragraph("GLOBAL ENERGY SOLUTIONS INC.")
    doc.add_paragraph("By: Maria Gonzales")
    doc.add_paragraph("Title: Chief Technology Officer")
    doc.add_paragraph("Date: May 18, 2024")

    doc.save("05_license_agreement_newyork_oil_gas.docx")
    print("✓ Created: 05_license_agreement_newyork_oil_gas.docx")


def create_employment_agreement_pdf():
    """Create Employment Agreement in PDF format"""
    filename = "06_employment_agreement_uk_tech.pdf"
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    Story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=16,
        textColor="black",
        spaceAfter=30,
        alignment=1,
    )

    Story.append(Paragraph("EMPLOYMENT AGREEMENT", title_style))
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph(
            'This Employment Agreement ("Agreement") is made on the 1st day of June, 2024, between:',
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph(
            "<b>EMPLOYER:</b> TechBridge Solutions Ltd, a company registered in England and Wales under company number 87654321, "
            'having its registered office at Innovation House, 25 Tech Street, Cambridge CB2 1AB, United Kingdom ("Company")',
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>AND</b>", styles["BodyText"]))
    Story.append(Spacer(1, 12))

    Story.append(
        Paragraph(
            '<b>EMPLOYEE:</b> Alexander Morrison, residing at 15 Park Lane, Cambridge CB1 2XY, United Kingdom ("Employee")',
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>1. POSITION AND DUTIES</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "The Company hereby employs the Employee as Senior Software Engineer, and the Employee accepts such employment upon "
            "the terms and conditions set out in this Agreement. The Employee shall perform duties related to software development, "
            "particularly in artificial intelligence and machine learning technologies, and shall report to the Chief Technology Officer.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>2. COMMENCEMENT AND TERM</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "This Agreement shall commence on 1st July 2024 and shall continue unless terminated by either party in accordance "
            "with the provisions of this Agreement.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>3. SALARY AND BENEFITS</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "3.1 Base Salary: The Company shall pay the Employee an annual salary of £85,000 (Eighty-Five Thousand British Pounds) "
            "payable in equal monthly installments in arrears on the last working day of each month.",
            styles["BodyText"],
        )
    )
    Story.append(
        Paragraph(
            "3.2 Annual Bonus: The Employee shall be eligible for an annual performance bonus of up to 20% of base salary, "
            "subject to achievement of agreed objectives.",
            styles["BodyText"],
        )
    )
    Story.append(
        Paragraph(
            "3.3 Benefits: The Employee shall be entitled to private medical insurance, life insurance, and 25 days of annual leave "
            "plus UK bank holidays.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>4. INTELLECTUAL PROPERTY</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "All inventions, discoveries, designs, and works created by the Employee during employment relating to the Company's "
            "business in the technology sector shall be the absolute property of the Company.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>5. CONFIDENTIALITY</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "The Employee agrees to maintain strict confidentiality regarding the Company's proprietary information, including "
            "source code, algorithms, customer data, and business strategies, both during employment and for a period of two (2) years "
            "following termination.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("<b>6. GOVERNING LAW</b>", styles["Heading2"]))
    Story.append(
        Paragraph(
            "This Agreement shall be governed by and construed in accordance with the laws of England and Wales, and the parties "
            "submit to the exclusive jurisdiction of the English courts.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 24))

    Story.append(
        Paragraph(
            "<b>IN WITNESS WHEREOF</b> the parties have executed this Agreement as of the date first above written.",
            styles["BodyText"],
        )
    )
    Story.append(Spacer(1, 24))

    Story.append(Paragraph("TECHBRIDGE SOLUTIONS LTD", styles["BodyText"]))
    Story.append(Paragraph("By: Rachel Thompson", styles["BodyText"]))
    Story.append(Paragraph("Title: Human Resources Director", styles["BodyText"]))
    Story.append(Paragraph("Date: 1st June 2024", styles["BodyText"]))
    Story.append(Spacer(1, 12))

    Story.append(Paragraph("EMPLOYEE", styles["BodyText"]))
    Story.append(Paragraph("Alexander Morrison", styles["BodyText"]))
    Story.append(Paragraph("Date: 1st June 2024", styles["BodyText"]))

    doc.build(Story)
    print("✓ Created: 06_employment_agreement_uk_tech.pdf")


def create_lease_agreement_docx():
    """Create Lease Agreement in DOCX format"""
    doc = Document()

    title = doc.add_heading("COMMERCIAL LEASE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'This Commercial Lease Agreement ("Lease") is entered into on 20th August 2024, between:'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("LANDLORD: ").bold = True
    p.add_run(
        "Emirates Property Holdings LLC, a limited liability company incorporated under the laws of the Emirate of Dubai, "
        'United Arab Emirates, with its registered office at Business Bay, Dubai, UAE ("Landlord")'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("AND").bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("TENANT: ").bold = True
    p.add_run(
        "Middle East Retail Group FZ-LLC, a free zone company incorporated in Dubai, UAE, with its office at Dubai Marina, "
        'Dubai, UAE ("Tenant")'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("PREMISES: ").bold = True
    p.add_run(
        "Commercial retail space located at Shop No. 15, Ground Floor, Dubai Mall Extension, Downtown Dubai, Dubai, "
        'United Arab Emirates, comprising approximately 2,500 square feet (the "Premises").'
    )

    doc.add_heading("1. LEASE GRANT", 2)
    doc.add_paragraph(
        "The Landlord hereby leases to the Tenant, and the Tenant hereby leases from the Landlord, the Premises for the "
        "term and upon the conditions hereinafter set forth."
    )

    doc.add_heading("2. TERM", 2)
    doc.add_paragraph(
        "The term of this Lease shall be for five (5) years, commencing on 1st October 2024 and ending on 30th September 2029 "
        '(the "Term"), unless sooner terminated as provided herein.'
    )

    doc.add_heading("3. RENT", 2)
    doc.add_paragraph(
        "3.1 Annual Rent: The Tenant shall pay annual rent of AED 1,500,000 (One Million Five Hundred Thousand UAE Dirhams) "
        "payable in four (4) equal quarterly installments of AED 375,000 each."
    )
    doc.add_paragraph(
        "3.2 Payment Terms: Rent shall be paid in advance on the first day of each quarter to the Landlord's designated "
        "bank account."
    )
    doc.add_paragraph(
        "3.3 Rent Increase: The annual rent shall increase by 5% at the beginning of each subsequent year of the Term."
    )

    doc.add_heading("4. SECURITY DEPOSIT", 2)
    doc.add_paragraph(
        "The Tenant shall pay a security deposit of AED 375,000 (Three Hundred Seventy-Five Thousand UAE Dirhams) upon "
        "execution of this Lease, to be held by the Landlord as security for the Tenant's performance."
    )

    doc.add_heading("5. USE OF PREMISES", 2)
    doc.add_paragraph(
        "The Premises shall be used solely for retail business operations in the fashion and apparel industry. The Tenant "
        "shall not use the Premises for any illegal purpose or in any manner that violates Dubai Municipality regulations."
    )

    doc.add_heading("6. INSURANCE", 2)
    doc.add_paragraph(
        "The Tenant shall maintain comprehensive general liability insurance with minimum coverage of AED 5,000,000 and "
        "contents insurance covering the full value of the Tenant's property and improvements."
    )

    doc.add_heading("7. COMPLIANCE WITH LAWS", 2)
    doc.add_paragraph(
        "The Tenant shall comply with all applicable laws, regulations, and ordinances of the Dubai Government, including "
        "but not limited to Dubai Municipality regulations, Dubai Civil Defence requirements, and Dubai Economic Department "
        "licensing requirements."
    )

    doc.add_heading("8. GOVERNING LAW AND JURISDICTION", 2)
    doc.add_paragraph(
        "This Lease shall be governed by and construed in accordance with the laws of the Emirate of Dubai, United Arab Emirates. "
        "Any disputes arising under this Lease shall be subject to the exclusive jurisdiction of the Dubai Courts or, at the "
        "parties' mutual agreement, the Dubai International Arbitration Centre (DIAC)."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("IN WITNESS WHEREOF, ").bold = True
    p.add_run(
        "the parties have executed this Commercial Lease Agreement as of the date first above written."
    )

    doc.add_paragraph()
    doc.add_paragraph("EMIRATES PROPERTY HOLDINGS LLC")
    doc.add_paragraph("By: Mohammed Al-Rashid")
    doc.add_paragraph("Title: Managing Director")
    doc.add_paragraph("Date: 20th August 2024")

    doc.add_paragraph()
    doc.add_paragraph("MIDDLE EAST RETAIL GROUP FZ-LLC")
    doc.add_paragraph("By: Sophie Laurent")
    doc.add_paragraph("Title: Chief Operating Officer")
    doc.add_paragraph("Date: 20th August 2024")

    doc.save("07_lease_agreement_dubai_real_estate.docx")
    print("✓ Created: 07_lease_agreement_dubai_real_estate.docx")


if __name__ == "__main__":
    import os

    os.chdir("/Users/jorgenino/Documents/legal_intel_dashboard/test_documents")

    print("Generating additional test documents...")
    print()

    create_franchise_agreement_pdf()
    create_license_agreement_docx()
    create_employment_agreement_pdf()
    create_lease_agreement_docx()

    print()
    print("✅ All additional test documents generated successfully!")
    print()
    print("Documents created:")
    print(
        "  4. 04_franchise_agreement_california.pdf - California Restaurant Franchise"
    )
    print("  5. 05_license_agreement_newyork_oil_gas.docx - New York Oil & Gas License")
    print("  6. 06_employment_agreement_uk_tech.pdf - UK Tech Employment")
    print("  7. 07_lease_agreement_dubai_real_estate.docx - Dubai Real Estate Lease")
